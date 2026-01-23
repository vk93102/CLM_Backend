import uuid
import hashlib
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Any
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from django.utils import timezone
import io
import os
import requests
from pathlib import Path
from PyPDF2 import PdfReader, PdfWriter
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.pdfgen import canvas
import logging

from .models import (
    Contract, ContractVersion, ContractClause, 
    ContractTemplate, Clause, BusinessRule, WorkflowLog,
    GenerationJob
)

logger = logging.getLogger(__name__)

# SignNow API Configuration
SIGNNOW_API_BASE = "https://api.signnow.com"
SIGNNOW_AUTH_URL = "https://app.signnow.com/oauth2/authorize"
SIGNNOW_TOKEN_URL = "https://api.signnow.com/oauth2/token"


class RuleEngine:
    @staticmethod
    def evaluate_condition(condition: Dict, context: Dict) -> bool:
        for key, expected_value in condition.items():
            if '__' in key:
                field, operator = key.rsplit('__', 1)
                actual_value = context.get(field)
                
                if actual_value is None:
                    return False
                
                if operator == 'gte' and not (actual_value >= expected_value):
                    return False
                elif operator == 'lte' and not (actual_value <= expected_value):
                    return False
                elif operator == 'gt' and not (actual_value > expected_value):
                    return False
                elif operator == 'lt' and not (actual_value < expected_value):
                    return False
                elif operator == 'in' and actual_value not in expected_value:
                    return False
                elif operator == 'contains' and expected_value not in actual_value:
                    return False
            else:
                if context.get(key) != expected_value:
                    return False
        
        return True
    
    @classmethod
    def get_mandatory_clauses(cls, tenant_id: uuid.UUID, contract_type: str, context: Dict) -> List[Dict]:
        rules = BusinessRule.objects.filter(
            tenant_id=tenant_id,
            rule_type='mandatory_clause',
            is_active=True
        ).filter(
            models.Q(contract_types=[]) | models.Q(contract_types__contains=[contract_type])
        ).order_by('-priority')
        
        mandatory_clauses = []
        for rule in rules:
            if cls.evaluate_condition(rule.conditions, context):
                mandatory_clauses.append({
                    'clause_id': rule.action.get('clause_id'),
                    'message': rule.action.get('message', ''),
                    'rule_name': rule.name
                })
        
        return mandatory_clauses
    
    @classmethod
    def get_clause_suggestions(cls, tenant_id: uuid.UUID, contract_type: str, context: Dict, clause_id: str) -> List[Dict]:
        try:
            clause = Clause.objects.get(
                tenant_id=tenant_id,
                clause_id=clause_id,
                status='published'
            )
        except Clause.DoesNotExist:
            return []
        
        suggestions = []
        
        for alt in clause.alternatives:
            trigger_rules = alt.get('trigger_rules', {})
            if not trigger_rules or cls.evaluate_condition(trigger_rules, context):
                suggestions.append({
                    'clause_id': alt['clause_id'],
                    'rationale': alt.get('rationale', 'Alternative clause'),
                    'confidence': alt.get('confidence', 0.8),
                    'source': 'predefined'
                })
        
        rules = BusinessRule.objects.filter(
            tenant_id=tenant_id,
            rule_type='clause_suggestion',
            is_active=True,
            action__target_clause_id=clause_id
        ).order_by('-priority')
        
        for rule in rules:
            if cls.evaluate_condition(rule.conditions, context):
                suggestions.append({
                    'clause_id': rule.action.get('suggest_clause_id'),
                    'rationale': rule.action.get('rationale', rule.description),
                    'confidence': rule.action.get('confidence', 0.7),
                    'source': 'rule_based'
                })
        
        return suggestions
    
    @classmethod
    def validate_contract(cls, tenant_id: uuid.UUID, contract_type: str, context: Dict, selected_clauses: List[str]) -> Tuple[bool, List[str]]:
        """
        Validate contract against business rules
        Returns: (is_valid, error_messages)
        """
        errors = []
        
        mandatory = cls.get_mandatory_clauses(tenant_id, contract_type, context)
        for req in mandatory:
            if req['clause_id'] not in selected_clauses:
                errors.append(f"Mandatory clause missing: {req['clause_id']} - {req['message']}")
        
        rules = BusinessRule.objects.filter(
            tenant_id=tenant_id,
            rule_type='validation',
            is_active=True
        ).filter(
            models.Q(contract_types=[]) | models.Q(contract_types__contains=[contract_type])
        )
        
        for rule in rules:
            if cls.evaluate_condition(rule.conditions, context):
                if rule.action.get('type') == 'error':
                    errors.append(rule.action.get('message', rule.description))
        
        return len(errors) == 0, errors


class ContractGenerator:
    def __init__(self, user_id: uuid.UUID, tenant_id: uuid.UUID):
        self.user_id = user_id
        self.tenant_id = tenant_id
        self.rule_engine = RuleEngine()
    
    def generate_from_template(
        self, 
        template_id: uuid.UUID,
        structured_inputs: Dict,
        user_instructions: Optional[str] = None,
        title: Optional[str] = None
    ) -> Contract:
        template = ContractTemplate.objects.get(
            id=template_id,
            tenant_id=self.tenant_id,
            status='published'
        )
        
        # Create contract record
        contract = Contract.objects.create(
            tenant_id=self.tenant_id,
            template=template,
            title=title or f"{template.contract_type} - {structured_inputs.get('counterparty', 'Draft')}",
            contract_type=template.contract_type,
            counterparty=structured_inputs.get('counterparty'),
            value=structured_inputs.get('value'),
            start_date=structured_inputs.get('start_date'),
            end_date=structured_inputs.get('end_date'),
            form_inputs=structured_inputs,
            user_instructions=user_instructions,
            created_by=self.user_id,
            status='draft',
            current_version=1
        )
        
        WorkflowLog.objects.create(
            contract=contract,
            action='created',
            performed_by=self.user_id,
            comment=f'Generated from template: {template.name} v{template.version}'
        )
        
        return contract
    
    def create_version(
        self, 
        contract: Contract, 
        selected_clauses: Optional[List[str]] = None,
        change_summary: Optional[str] = None
    ) -> ContractVersion:
        template = contract.template
        context = {
            'contract_type': contract.contract_type,
            'contract_value': float(contract.value) if contract.value else 0,
            'counterparty': contract.counterparty,
            **contract.form_inputs
        }
        
        if selected_clauses is None:
            selected_clauses = list(template.mandatory_clauses)
            mandatory = self.rule_engine.get_mandatory_clauses(
                self.tenant_id, 
                contract.contract_type, 
                context
            )
            for req in mandatory:
                if req['clause_id'] not in selected_clauses:
                    selected_clauses.append(req['clause_id'])
        
        is_valid, errors = self.rule_engine.validate_contract(
            self.tenant_id,
            contract.contract_type,
            context,
            selected_clauses
        )
        
        if not is_valid:
            raise ValidationError({"clauses": errors})
        
        doc = self._create_document(contract, selected_clauses, context)
        
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        content = doc_bytes.getvalue()
        file_hash = hashlib.sha256(content).hexdigest()
        
        version_number = contract.current_version
        version = ContractVersion.objects.create(
            contract=contract,
            version_number=version_number,
            template_id=template.id,
            template_version=template.version,
            change_summary=change_summary or f'Version {version_number}',
            created_by=self.user_id,
            file_size=len(content),
            file_hash=file_hash,
            r2_key=f'contracts/{contract.id}/v{version_number}.docx'
        )
        
        self._store_clause_provenance(version, selected_clauses, context)
        
        contract.current_version = version_number + 1
        contract.save()
        
        WorkflowLog.objects.create(
            contract=contract,
            action='version_created',
            performed_by=self.user_id,
            comment=f'Version {version_number} created',
            metadata={'clause_count': len(selected_clauses)}
        )
        
        return version
    
    def _create_document(self, contract: Contract, clause_ids: List[str], context: Dict) -> Document:
        doc = Document()
        doc.add_heading(contract.title, 0)
        p = doc.add_paragraph()
        p.add_run(f"Contract Type: ").bold = True
        p.add_run(f"{contract.contract_type}\n")
        p.add_run(f"Counterparty: ").bold = True
        p.add_run(f"{contract.counterparty or 'N/A'}\n")
        if contract.value:
            p.add_run(f"Value: ").bold = True
            p.add_run(f"${contract.value:,.2f}\n")
        p.add_run(f"Created: ").bold = True
        p.add_run(f"{contract.created_at.strftime('%Y-%m-%d')}\n")
        doc.add_paragraph()  
        
        clauses = Clause.objects.filter(
            tenant_id=self.tenant_id,
            clause_id__in=clause_ids,
            status='published'
        ).order_by('clause_id')
        
        for i, clause in enumerate(clauses, 1):
            heading = doc.add_heading(f"{i}. {clause.name}", level=2)
            content = self._replace_merge_fields(clause.content, context)
            doc.add_paragraph(content)
            provenance = doc.add_paragraph()
            provenance.add_run(f"[Clause ID: {clause.clause_id} v{clause.version}]").font.size = Pt(8)
            provenance.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            doc.add_paragraph() 
        
        return doc
    
    def _replace_merge_fields(self, text: str, context: Dict) -> str:
        result = text
        for key, value in context.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in result:
                result = result.replace(placeholder, str(value))
        return result
    
    def _store_clause_provenance(self, version: ContractVersion, clause_ids: List[str], context: Dict):
        clauses = Clause.objects.filter(
            tenant_id=self.tenant_id,
            clause_id__in=clause_ids,
            status='published'
        ).order_by('clause_id')
        
        for position, clause in enumerate(clauses, 1):
            alternatives = self.rule_engine.get_clause_suggestions(
                self.tenant_id,
                version.contract.contract_type,
                context,
                clause.clause_id
            )
            
            ContractClause.objects.create(
                contract_version=version,
                clause_id=clause.clause_id,
                clause_version=clause.version,
                clause_name=clause.name,
                clause_content=clause.content,
                is_mandatory=clause.is_mandatory,
                position=position,
                alternatives_suggested=alternatives
            )


from rest_framework.exceptions import ValidationError
from django.db import models



# ========== GENERATION SERVICE ==========

class ContractGenerationService:
    """
    Service for generating contracts from templates and clauses
    """
    
    def __init__(self, tenant_id: str, user_id: str):
        self.tenant_id = tenant_id
        self.user_id = user_id
    
    def generate_contract(
        self,
        template_id: str,
        form_inputs: Dict[str, Any],
        user_instructions: Optional[str] = None,
        async_mode: bool = False
    ) -> Contract:
        """
        Generate a contract from template and inputs
        
        Args:
            template_id: UUID of the template to use
            form_inputs: Structured form data for merge fields
            user_instructions: Optional free-text instructions
            async_mode: Whether to generate asynchronously
            
        Returns:
            Contract instance
        """
        # Load template
        template = ContractTemplate.objects.get(
            id=template_id,
            tenant_id=self.tenant_id,
            status='published'
        )
        
        # Validate business rules
        validation_errors = self._validate_business_rules(template, form_inputs)
        if validation_errors:
            raise ValueError(f"Business rule violations: {validation_errors}")
        
        # Validate mandatory clauses
        missing_clauses = self._check_mandatory_clauses(template)
        if missing_clauses:
            raise ValueError(f"Missing mandatory clauses: {missing_clauses}")
        
        # Create contract record
        contract = Contract.objects.create(
            tenant_id=self.tenant_id,
            template=template,
            title=form_inputs.get('contract_title', 'Untitled Contract'),
            contract_type=template.contract_type,
            created_by=self.user_id,
            counterparty=form_inputs.get('counterparty'),
            value=form_inputs.get('value'),
            start_date=form_inputs.get('start_date'),
            end_date=form_inputs.get('end_date'),
            form_inputs=form_inputs,
            user_instructions=user_instructions,
            status='draft',
            current_version=1
        )
        
        if async_mode:
            # Create async job
            job = GenerationJob.objects.create(
                contract=contract,
                status='pending'
            )
            # TODO: Queue celery task
            return contract
        else:
            # Generate synchronously
            self._generate_document(contract, template, form_inputs)
            return contract
    
    def _generate_document(
        self,
        contract: Contract,
        template: ContractTemplate,
        form_inputs: Dict[str, Any]
    ) -> ContractVersion:
        """
        Generate the actual DOCX document
        """
        # Create new document (in production, load template from R2)
        doc = Document()
        
        # Add title
        title = doc.add_heading(form_inputs.get('contract_title', 'Contract'), 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Contract Type: {template.contract_type}")
        doc.add_paragraph(f"Version: 1")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")
        
        # Populate merge fields
        for field, value in form_inputs.items():
            if value:
                doc.add_paragraph(f"{field.replace('_', ' ').title()}: {value}")
        
        doc.add_paragraph("")
        
        # Assemble clauses
        clauses = self._get_applicable_clauses(template)
        
        doc.add_heading("Contract Clauses", 1)
        
        clause_records = []
        for idx, clause in enumerate(clauses, 1):
            # Add clause
            heading = doc.add_heading(f"{idx}. {clause.name}", 2)
            
            # Add clause metadata (for provenance)
            metadata = doc.add_paragraph(
                f"[Clause ID: {clause.clause_id} | Version: {clause.version}]",
                style='Intense Quote'
            )
            metadata_run = metadata.runs[0]
            metadata_run.font.size = Pt(8)
            metadata_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Add clause content
            doc.add_paragraph(clause.content)
            doc.add_paragraph("")
            
            # Track clause usage
            clause_records.append({
                'clause': clause,
                'position': idx,
                'alternatives': self._get_clause_alternatives(clause)
            })
        
        # Save document to memory
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        # Calculate hash
        doc_content = doc_buffer.getvalue()
        file_hash = hashlib.sha256(doc_content).hexdigest()
        
        # In production: Upload to R2
        r2_key = f"contracts/{contract.tenant_id}/{contract.id}/v1.docx"
        # r2_service.upload(r2_key, doc_buffer)
        
        # Create contract version
        version = ContractVersion.objects.create(
            contract=contract,
            version_number=1,
            r2_key=r2_key,
            template_id=template.id,
            template_version=template.version,
            change_summary="Initial generation",
            created_by=self.user_id,
            file_size=len(doc_content),
            file_hash=file_hash
        )
        
        # Create clause provenance records
        for record in clause_records:
            clause = record['clause']
            ContractClause.objects.create(
                contract_version=version,
                clause_id=clause.clause_id,
                clause_version=clause.version,
                clause_name=clause.name,
                clause_content=clause.content,
                is_mandatory=clause.is_mandatory,
                position=record['position'],
                alternatives_suggested=record['alternatives']
            )
        
        return version
    
    def _validate_business_rules(
        self,
        template: ContractTemplate,
        form_inputs: Dict[str, Any]
    ) -> List[str]:
        """
        Validate business rules from template
        """
        errors = []
        rules = template.business_rules
        
        # Check required fields
        required_fields = rules.get('required_fields', [])
        for field in required_fields:
            if field not in form_inputs or not form_inputs[field]:
                errors.append(f"Missing required field: {field}")
        
        # Check value constraints
        if 'min_value' in rules and 'value' in form_inputs:
            if float(form_inputs['value']) < float(rules['min_value']):
                errors.append(f"Value must be at least {rules['min_value']}")
        
        if 'max_value' in rules and 'value' in form_inputs:
            if float(form_inputs['value']) > float(rules['max_value']):
                errors.append(f"Value must not exceed {rules['max_value']}")
        
        # Check jurisdiction (if required)
        if 'allowed_jurisdictions' in rules and 'jurisdiction' in form_inputs:
            if form_inputs['jurisdiction'] not in rules['allowed_jurisdictions']:
                errors.append(
                    f"Jurisdiction must be one of: {', '.join(rules['allowed_jurisdictions'])}"
                )
        
        return errors
    
    def _check_mandatory_clauses(self, template: ContractTemplate) -> List[str]:
        """
        Check if all mandatory clauses exist
        """
        mandatory_clause_ids = template.mandatory_clauses
        missing = []
        
        for clause_id in mandatory_clause_ids:
            exists = Clause.objects.filter(
                tenant_id=self.tenant_id,
                clause_id=clause_id,
                status='published'
            ).exists()
            
            if not exists:
                missing.append(clause_id)
        
        return missing
    
    def _get_applicable_clauses(self, template: ContractTemplate) -> List[Clause]:
        """
        Get all applicable clauses for this template
        """
        # Get mandatory clauses first
        mandatory_clauses = list(Clause.objects.filter(
            tenant_id=self.tenant_id,
            clause_id__in=template.mandatory_clauses,
            contract_type=template.contract_type,
            status='published'
        ).order_by('clause_id'))
        
        # Get optional clauses
        optional_clauses = list(Clause.objects.filter(
            tenant_id=self.tenant_id,
            contract_type=template.contract_type,
            is_mandatory=False,
            status='published'
        ).exclude(
            clause_id__in=template.mandatory_clauses
        ).order_by('clause_id'))
        
        # Combine (mandatory first)
        return mandatory_clauses + optional_clauses
    
    def _get_clause_alternatives(self, clause: Clause) -> List[Dict[str, Any]]:
        """
        Get alternative clauses with confidence scores
        """
        alternatives = []
        
        for alt_data in clause.alternatives:
            alt_clause_id = alt_data.get('clause_id')
            rationale = alt_data.get('rationale', '')
            confidence = alt_data.get('confidence', 0.5)
            
            # Try to get the alternative clause
            try:
                alt_clause = Clause.objects.get(
                    tenant_id=self.tenant_id,
                    clause_id=alt_clause_id,
                    status='published'
                )
                
                alternatives.append({
                    'clause_id': alt_clause_id,
                    'clause_name': alt_clause.name,
                    'rationale': rationale,
                    'confidence': confidence
                })
            except Clause.DoesNotExist:
                continue
        
        return alternatives
    
    def approve_contract(self, contract_id: str) -> Contract:
        """
        Approve a contract for download/send/sign
        """
        contract = Contract.objects.get(
            id=contract_id,
            tenant_id=self.tenant_id
        )
        
        contract.is_approved = True
        contract.approved_by = self.user_id
        contract.approved_at = datetime.now()
        contract.status = 'approved'
        contract.save()
        
        return contract
    
    def create_new_version(
        self,
        contract_id: str,
        changes: Dict[str, Any],
        change_summary: str
    ) -> ContractVersion:
        """
        Create a new contract version (never overwrite)
        """
        contract = Contract.objects.get(
            id=contract_id,
            tenant_id=self.tenant_id
        )
        
        # Increment version
        new_version_number = contract.current_version + 1
        
        # Merge changes with existing inputs
        updated_inputs = {**contract.form_inputs, **changes}
        
        # Update contract
        contract.form_inputs = updated_inputs
        contract.current_version = new_version_number
        contract.is_approved = False  # Reset approval
        contract.save()
        
        # Regenerate document
        template = contract.template
        version = self._generate_document(contract, template, updated_inputs)
        version.version_number = new_version_number
        version.change_summary = change_summary
        version.save()
        
        return version

# ========== CONTRACT GENERATOR SERVICE ==========

# Template paths mapping
TEMPLATE_PATHS = {
    'nda': 'templates/nda_template.pdf',
    'agency_agreement': 'templates/agency_agreement_template.pdf',
    'property_management': 'templates/property_management_template.pdf',
    'employment_contract': 'templates/employment_contract_template.pdf',
}

# Contract field definitions
CONTRACT_FIELD_DEFINITIONS = {
    'nda': {
        'required': ['date', '1st_party_name', '2nd_party_name', 'agreement_type'],
        'optional': ['1st_party_relationship', '2nd_party_relationship', 'governing_law',
                    '1st_party_printed_name', '1st_party_signature', 'signature_date_1st_party',
                    '2nd_party_printed_name', '2nd_party_signature', 'signature_date_2nd_party']
    },
    'agency_agreement': {
        'required': ['effective_date', 'principal_name', 'agent_name', 'monthly_compensation'],
        'optional': ['principal_address', 'agent_address', 'invoice_submission_date', 
                    'payment_method', 'payment_address', 'contract_end_date']
    },
    'property_management': {
        'required': ['effective_date', 'owner_name', 'manager_name', 'property_address'],
        'optional': ['owner_address', 'manager_address', 'repair_approval_limit', 'governing_law']
    },
    'employment_contract': {
        'required': ['effective_date', 'employer_name', 'employee_name', 'job_title'],
        'optional': ['employer_address', 'employee_address', 'commencement_date', 
                    'employment_type', 'salary_amount', 'payment_schedule']
    }
}

class ContractGeneratorService:
    """Service for generating filled PDF contracts"""
    
    def __init__(self, base_path: str = '/Users/vishaljha/CLM_Backend'):
        """Initialize the service with base path"""
        self.base_path = Path(base_path)
        self.templates_path = self.base_path / 'templates'
        
    def validate_contract_type(self, contract_type: str) -> bool:
        """Validate if contract type is supported"""
        return contract_type.lower() in TEMPLATE_PATHS
    
    def get_required_fields(self, contract_type: str) -> Dict[str, list]:
        """Get required and optional fields for a contract type"""
        contract_type = contract_type.lower()
        if contract_type not in CONTRACT_FIELD_DEFINITIONS:
            raise ValueError(f"Unknown contract type: {contract_type}")
        return CONTRACT_FIELD_DEFINITIONS[contract_type]
    
    def validate_contract_data(self, contract_type: str, data: Dict[str, Any]) -> tuple[bool, list]:
        """
        Validate that all required fields are present
        Returns: (is_valid, missing_fields)
        """
        contract_type = contract_type.lower()
        fields_config = self.get_required_fields(contract_type)
        required_fields = fields_config.get('required', [])
        
        missing_fields = []
        for field in required_fields:
            if field not in data or not str(data[field]).strip():
                missing_fields.append(field)
        
        is_valid = len(missing_fields) == 0
        return is_valid, missing_fields
    
    def generate_contract(
        self,
        contract_type: str,
        data: Dict[str, Any],
        output_path: Optional[str] = None,
    ) -> str:
        """
        Generate a filled contract PDF
        
        Args:
            contract_type: Type of contract ('nda', 'agency_agreement', etc.)
            data: Dictionary with contract field values
            output_path: Optional path to save the PDF. If None, generates in temp location
        
        Returns:
            Path to the generated PDF file
        """
        # Validate contract type
        contract_type = contract_type.lower()
        if not self.validate_contract_type(contract_type):
            raise ValueError(f"Unsupported contract type: {contract_type}")
        
        # Validate data
        is_valid, missing_fields = self.validate_contract_data(contract_type, data)
        if not is_valid:
            raise ValueError(f"Missing required fields: {', '.join(missing_fields)}")
        
        # Get template path
        template_relative = TEMPLATE_PATHS[contract_type]
        template_path = self.base_path / template_relative
        
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        logger.info(f"Generating {contract_type} contract with data keys: {list(data.keys())}")
        
        # Create filled PDF
        filled_pdf = self._fill_pdf_template(
            str(template_path),
            contract_type,
            data
        )
        
        # Save to disk
        if output_path is None:
            output_path = self.base_path / f"generated_{contract_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        with open(output_path, 'wb') as f:
            f.write(filled_pdf)
        
        logger.info(f"Contract generated successfully: {output_path}")
        return str(output_path)
    
    def _fill_pdf_template(
        self,
        template_path: str,
        contract_type: str,
        data: Dict[str, Any]
    ) -> bytes:
        """
        Fill PDF template with data and return as bytes
        Uses ReportLab to overlay text on PDF
        """
        # Read the template PDF
        with open(template_path, 'rb') as f:
            reader = PdfReader(f)
            writer = PdfWriter()
            
            # Get the first page (assuming single page templates)
            page = reader.pages[0]
            
            # Create an overlay with filled data
            overlay_stream = io.BytesIO()
            overlay_canvas = canvas.Canvas(overlay_stream, pagesize=letter)
            
            # Define positioning based on contract type
            positions = self._get_field_positions(contract_type)
            
            # Fill in the fields
            for field_name, value in data.items():
                if field_name in positions:
                    x, y = positions[field_name]
                    self._draw_field_value(overlay_canvas, field_name, str(value), x, y, contract_type)
            
            overlay_canvas.save()
            overlay_stream.seek(0)
            
            # Read the overlay
            overlay_reader = PdfReader(overlay_stream)
            overlay_page = overlay_reader.pages[0]
            
            # Merge overlay with template
            page.merge_page(overlay_page)
            writer.add_page(page)
            
            # Write output
            output_stream = io.BytesIO()
            writer.write(output_stream)
            output_stream.seek(0)
            return output_stream.read()
    
    def _get_field_positions(self, contract_type: str) -> Dict[str, tuple]:
        """
        Get (x, y) positions for fields on each contract type
        Positions are approximate based on standard PDF layouts
        """
        if contract_type == 'nda':
            return {
                'date': (450, 750),
                '1st_party_name': (250, 700),
                '2nd_party_name': (250, 670),
                'agreement_type': (120, 640),
                '1st_party_relationship': (250, 600),
                '2nd_party_relationship': (250, 570),
                'governing_law': (250, 540),
                '1st_party_printed_name': (100, 200),
                '1st_party_signature': (100, 170),
                'signature_date_1st_party': (300, 170),
                '2nd_party_printed_name': (400, 200),
                '2nd_party_signature': (400, 170),
                'signature_date_2nd_party': (600, 170),
            }
        
        elif contract_type == 'agency_agreement':
            return {
                'effective_date': (450, 750),
                'principal_name': (250, 710),
                'principal_address': (250, 680),
                'agent_name': (250, 640),
                'agent_address': (250, 610),
                'monthly_compensation': (250, 520),
                'invoice_submission_date': (250, 490),
                'payment_method': (250, 460),
                'payment_address': (250, 430),
                'contract_end_date': (250, 400),
                'principal_printed_name': (100, 150),
                'principal_signature': (100, 120),
                'principal_signature_date': (300, 120),
                'agent_printed_name': (400, 150),
                'agent_signature': (400, 120),
                'agent_signature_date': (600, 120),
            }
        
        elif contract_type == 'property_management':
            return {
                'effective_date': (450, 750),
                'owner_name': (250, 710),
                'owner_address': (250, 680),
                'manager_name': (250, 640),
                'manager_address': (250, 610),
                'property_address': (250, 560),
                'repair_approval_limit': (250, 530),
                'governing_law': (250, 500),
                'owner_printed_name': (100, 150),
                'owner_signature': (100, 120),
                'owner_signature_date': (300, 120),
                'manager_printed_name': (400, 150),
                'manager_signature': (400, 120),
                'manager_signature_date': (600, 120),
            }
        
        elif contract_type == 'employment_contract':
            return {
                'effective_date': (450, 750),
                'employer_name': (250, 720),
                'employer_address': (250, 690),
                'employee_name': (250, 650),
                'employee_address': (250, 620),
                'job_title': (250, 580),
                'responsibilities': (250, 550),
                'commencement_date': (250, 510),
                'working_days_from': (150, 480),
                'working_days_to': (350, 480),
                'employment_type': (150, 450),
                'working_hours_from': (250, 420),
                'working_hours_to': (350, 420),
                'average_hours_per_week': (250, 390),
                'daily_lunch_break': (250, 360),
                'work_mode': (150, 330),
                'work_location': (250, 300),
                'salary_amount': (250, 260),
                'payment_schedule': (250, 230),
                'payment_method': (250, 200),
                'probation_period': (250, 170),
                'insurance_benefit': (250, 130),
                'holidays': (250, 100),
                'employer_printed_name': (100, 50),
                'employee_printed_name': (400, 50),
            }
        
        return {}
    
    def _draw_field_value(
        self,
        canvas_obj: canvas.Canvas,
        field_name: str,
        value: str,
        x: float,
        y: float,
        contract_type: str
    ) -> None:
        """Draw a field value on the canvas"""
        try:
            # Format value based on field type
            if 'date' in field_name and not value:
                value = datetime.now().strftime('%m/%d/%Y')
            
            # Truncate long values
            if len(value) > 50:
                value = value[:47] + "..."
            
            # Set font
            canvas_obj.setFont("Helvetica", 10)
            canvas_obj.setFillColor(colors.black)
            
            # Handle checkboxes
            if field_name == 'agreement_type':
                # This would be handled separately in a real implementation
                pass
            elif field_name == 'employment_type':
                # This would be handled separately
                pass
            elif field_name == 'work_mode':
                # This would be handled separately
                pass
            else:
                # Draw regular text
                canvas_obj.drawString(x, y, value)
            
            logger.debug(f"Drew {field_name} at ({x}, {y}): {value[:30]}")
        
        except Exception as e:
            logger.warning(f"Error drawing field {field_name}: {str(e)}")


# Convenience function
def generate_contract(
    contract_type: str,
    data: Dict[str, Any],
    output_path: Optional[str] = None,
) -> str:
    """
    Generate a contract - convenience function
    """
    service = ContractGeneratorService()
    return service.generate_contract(contract_type, data, output_path)

# ========== SIGNNOW SERVICE ==========

class SignNowAuthService:
    """
    Manages SignNow OAuth token lifecycle
    """
    
    CACHE_KEY_TOKEN = "signnow_access_token"
    CACHE_KEY_REFRESH = "signnow_refresh_token"
    
    def __init__(self):
        # Get credentials from database instead of environment variables
        from .models import SignNowCredential
        cred = SignNowCredential.objects.first()
        if cred:
            self.client_id = cred.client_id
            self.client_secret = cred.client_secret
            logger.info(f"Loaded SignNow credentials from database: {cred.client_id[:20]}...")
        else:
            # Fallback to environment variables if no database credential
            self.client_id = os.getenv("SIGNNOW_CLIENT_ID")
            self.client_secret = os.getenv("SIGNNOW_CLIENT_SECRET")
            logger.warning("No SignNow credential in database, using environment variables")
        self.redirect_uri = os.getenv("SIGNNOW_REDIRECT_URI", "http://localhost:11000/api/signnow/callback/")
    
    def get_access_token(self):
        """
        Get cached access token or refresh if expired
        Returns: (token, is_fresh)
        """
        # Try cache first
        token = cache.get(self.CACHE_KEY_TOKEN)
        if token:
            logger.info("Using cached SignNow access token")
            return token, False
        
        # Try to refresh
        refreshed_token = self.refresh_token()
        if refreshed_token:
            logger.info("Refreshed SignNow access token")
            return refreshed_token, True
        
        logger.error("Could not obtain SignNow access token")
        raise Exception("SignNow authentication failed - no valid token")
    
    def refresh_token(self):
        """
        Refresh OAuth token using refresh_token
        Returns: access_token or None
        """
        from .models import SignNowCredential
        
        try:
            cred = SignNowCredential.objects.first()
            if not cred or not cred.refresh_token:
                logger.error("No SignNow credential with refresh token")
                return None
            
            # Call SignNow token endpoint
            payload = {
                "grant_type": "refresh_token",
                "refresh_token": cred.refresh_token,
                "client_id": self.client_id,
                "client_secret": self.client_secret,
            }
            
            response = requests.post(
                SIGNNOW_TOKEN_URL,
                data=payload,  # Use 'data' for form encoding (not json)
                timeout=10
            )
            response.raise_for_status()
            
            data = response.json()
            access_token = data.get("access_token")
            new_refresh_token = data.get("refresh_token", cred.refresh_token)
            expires_in = data.get("expires_in", 3600)
            
            # Update credential
            cred.access_token = access_token
            cred.refresh_token = new_refresh_token
            cred.token_expires_at = timezone.now() + timedelta(seconds=expires_in)
            cred.save()
            
            # Cache token
            cache.set(self.CACHE_KEY_TOKEN, access_token, expires_in - 300)  # 5 min buffer
            
            logger.info("Successfully refreshed SignNow token")
            return access_token
            
        except Exception as e:
            logger.error(f"Token refresh failed: {str(e)}")
            return None
    
    def set_initial_token(self, access_token, refresh_token=None, expires_in=3600):
        """
        Store initial OAuth token after first authentication
        """
        from .models import SignNowCredential
        
        try:
            cred = SignNowCredential.objects.first()
            if not cred:
                logger.error("No SignNow credential to update")
                return False
            
            cred.access_token = access_token
            if refresh_token:
                cred.refresh_token = refresh_token
            cred.token_expires_at = timezone.now() + timedelta(seconds=expires_in)
            cred.save()
            
            # Cache token
            cache.set(self.CACHE_KEY_TOKEN, access_token, expires_in - 300)
            
            logger.info("Stored initial SignNow token")
            return True
            
        except Exception as e:
            logger.error(f"Failed to store initial token: {str(e)}")
            return False


class SignNowAPIService:
    """
    Handles all SignNow API operations
    """
    
    def __init__(self):
        self.auth_service = SignNowAuthService()
        self.base_url = SIGNNOW_API_BASE
    
    def _get_headers(self):
        """Get authorization headers with current access token"""
        token, _ = self.auth_service.get_access_token()
        return {
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json",
        }
    
    def _request(self, method, endpoint, **kwargs):
        """
        Make authenticated request to SignNow API
        Handles error responses and logging
        """
        url = f"{self.base_url}{endpoint}"
        headers = self._get_headers()
        
        if "headers" in kwargs:
            headers.update(kwargs.pop("headers"))
        
        try:
            response = requests.request(
                method,
                url,
                headers=headers,
                timeout=30,
                **kwargs
            )
            
            # Log request
            logger.info(f"SignNow API: {method} {endpoint} -> {response.status_code}")
            
            # Check for errors
            if response.status_code >= 400:
                logger.error(
                    f"SignNow API Error: {response.status_code}\n"
                    f"URL: {url}\n"
                    f"Response: {response.text}"
                )
                response.raise_for_status()
            
            return response
            
        except requests.exceptions.RequestException as e:
            logger.error(f"SignNow API request failed: {str(e)}")
            raise
    
    # ═══════════════════════════════════════════════════════════════════════════
    # DOCUMENT OPERATIONS
    # ═══════════════════════════════════════════════════════════════════════════
    
    def upload_document(self, pdf_file, document_name):
        """
        Upload PDF document to SignNow
        
        Args:
            pdf_file: Binary PDF content
            document_name: Name for the document in SignNow
        
        Returns:
            {
                'id': 'doc_id',
                'document_id': 'doc_id',
                'owner_id': 'owner_id',
                'pages': [...]
            }
        """
        files = {
            'file': (f"{document_name}.pdf", pdf_file, 'application/pdf')
        }
        
        response = self._request(
            "POST",
            "/v2/documents",
            files=files
        )
        
        data = response.json()
        logger.info(f"Uploaded document to SignNow: {data.get('id')}")
        return data
    
    def get_document_details(self, document_id):
        """
        Get document details from SignNow
        
        Returns:
            {
                'id': 'doc_id',
                'owner_id': 'owner_id',
                'pages': [...],
                'invitations': [...],
                'status': 'completed' | 'pending' | ...
            }
        """
        response = self._request(
            "GET",
            f"/v2/documents/{document_id}"
        )
        
        return response.json()
    
    def download_document(self, document_id):
        """
        Download executed PDF from SignNow
        
        Returns: Binary PDF content
        """
        response = self._request(
            "GET",
            f"/v2/documents/{document_id}/download"
        )
        
        return response.content
    
    # ═══════════════════════════════════════════════════════════════════════════
    # INVITATION OPERATIONS
    # ═══════════════════════════════════════════════════════════════════════════
    
    def create_invite(self, document_id, signers, signing_order="sequential"):
        """
        Create signature invitation(s) for document
        
        Args:
            document_id: SignNow document ID
            signers: List of signer dicts
                [
                    {"email": "signer@example.com", "name": "Signer Name", "order": 1},
                    ...
                ]
            signing_order: "sequential" or "parallel"
        
        Returns:
            Invitation response with invite IDs
        """
        payload = {
            "to": [
                {
                    "email": signer["email"],
                    "role_id": str(i + 1),  # SignNow role ID
                    "order": signer.get("order", i + 1) if signing_order == "sequential" else 0,
                    "name": signer.get("name", ""),
                }
                for i, signer in enumerate(signers)
            ],
            "from": "service_account@company.com",  # Service account sender
            "subject": "Please sign this document",
            "message": "Please review and sign the attached document.",
        }
        
        response = self._request(
            "POST",
            f"/v2/documents/{document_id}/invites",
            json=payload
        )
        
        data = response.json()
        logger.info(f"Created invitation for document {document_id}")
        return data
    
    def get_invite_status(self, invite_id):
        """
        Get status of a specific invitation
        
        Returns:
            {
                'id': 'invite_id',
                'status': 'pending' | 'viewed' | 'signed' | 'declined',
                'created': timestamp,
                'updated': timestamp
            }
        """
        response = self._request(
            "GET",
            f"/v2/invites/{invite_id}"
        )
        
        return response.json()
    
    # ═══════════════════════════════════════════════════════════════════════════
    # SIGNING URL OPERATIONS
    # ═══════════════════════════════════════════════════════════════════════════
    
    def get_signing_link(self, document_id, signer_email):
        """
        Generate embedded signing URL for a signer
        
        Args:
            document_id: SignNow document ID
            signer_email: Email of signer
        
        Returns:
            {
                'signing_link': 'https://...',
                'expires_at': timestamp
            }
        """
        payload = {
            "email": signer_email,
        }
        
        response = self._request(
            "POST",
            f"/v2/documents/{document_id}/signing-link",
            json=payload
        )
        
        data = response.json()
        logger.info(f"Generated signing link for {signer_email}")
        return data
    
    # ═══════════════════════════════════════════════════════════════════════════
    # STATUS POLLING
    # ═══════════════════════════════════════════════════════════════════════════
    
    def get_document_status(self, document_id):
        """
        Poll document status from SignNow
        
        Returns normalized status with detailed signer info
        """
        doc = self.get_document_details(document_id)
        
        # Normalize status
        signnow_status = doc.get("status", "unknown")
        
        # Map SignNow status to our status
        status_map = {
            "pending": "sent",
            "viewed": "in_progress",
            "in_progress": "in_progress",
            "completed": "completed",
            "declined": "declined",
            "expired": "expired",
        }
        
        normalized_status = status_map.get(signnow_status, signnow_status)
        
        # Extract signer statuses from invitations
        signers_status = []
        for invite in doc.get("invites", []):
            signers_status.append({
                "email": invite.get("email"),
                "invite_id": invite.get("id"),
                "status": invite.get("status"),
                "signed_at": invite.get("signed_at"),
            })
        
        return {
            "document_id": document_id,
            "status": normalized_status,
            "signnow_status": signnow_status,
            "signers": signers_status,
            "is_completed": normalized_status == "completed",
            "updated_at": datetime.now().isoformat(),
        }
